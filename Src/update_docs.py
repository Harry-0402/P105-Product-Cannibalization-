import os
import docx

def replace_text_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                # To preserve formatting, we replace text inside the runs
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

def update_document(file_path, replacements):
    try:
        doc = docx.Document(file_path)
        replace_text_in_paragraphs(doc.paragraphs, replacements)
        replace_text_in_tables(doc.tables, replacements)
        doc.save(file_path)
        print(f"Successfully updated: {file_path}")
    except Exception as e:
        print(f"Error updating {file_path}: {e}")

if __name__ == "__main__":
    directory = "."
    files = [
        "BRD_Product_Cannibalization_Analysis.docx",
        "Data_Dictionary.docx",
        "FRD_Product_Cannibalization_Analysis.docx",
        "PRD_Product_Cannibalization_Analysis.docx",
        "Project_Charter.docx",
        "Test_Plan_Validation_Document.docx",
        "User_Manual_Dashboard_Guide.docx"
    ]

    replacements = {
        "Power BI": "Interactive HTML/JS Dashboard",
        "PowerBI": "HTML Dashboard",
        "DAX": "JavaScript",
        "Power Query": "Pandas DataFrames",
        ".pbix": "index.html",
        "Power BI Service": "Web Browser",
        "Publish to Power BI": "Deploy HTML Dashboard",
        "Python scripts": "Jupyter Notebooks",
        "Python scripts (.py)": "Jupyter Notebooks (.ipynb)",
        ".py format": ".ipynb format",
        "Power BI dashboard": "Interactive HTML dashboard",
        "Power BI dashboards": "Interactive HTML dashboards"
    }

    for file in files:
        if os.path.exists(file):
            update_document(file, replacements)
        else:
            print(f"File not found: {file}")
