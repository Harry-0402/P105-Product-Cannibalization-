import os
import docx

def replace_text_in_paragraphs(paragraphs, replacements):
    for paragraph in paragraphs:
        for old_text, new_text in replacements.items():
            if old_text in paragraph.text:
                for run in paragraph.runs:
                    if old_text in run.text:
                        run.text = run.text.replace(old_text, new_text)

def replace_text_in_tables(tables, replacements):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_text_in_paragraphs(cell.paragraphs, replacements)

def update_document(file_path, replacements):
    try:
        doc = docx.Document(file_path)
        replace_text_in_paragraphs(doc.paragraphs, replacements)
        replace_text_in_tables(doc.tables, replacements)
        doc.save(file_path)
        print(f"Successfully cleaned: {file_path}")
    except Exception as e:
        print(f"Error cleaning {file_path}: {e}")

if __name__ == "__main__":
    files = [
        "BRD_Product_Cannibalization_Analysis.docx",
        "Data_Dictionary.docx",
        "FRD_Product_Cannibalization_Analysis.docx",
        "PRD_Product_Cannibalization_Analysis.docx",
        "Project_Charter.docx",
        "Test_Plan_Validation_Document.docx",
        "User_Manual_Dashboard_Guide.docx"
    ]

    replacements = {
        "Dashboard dashboard": "Dashboard",
        "Dashboard dashboards": "Dashboards",
        "Interactive HTML dashboards": "Interactive HTML Dashboards",
        "Interactive HTML dashboard": "Interactive HTML Dashboard",
        "Interactive HTML/JS Dashboard dashboard": "Interactive HTML/JS Dashboard",
        "Interactive HTML/JS Dashboard Dashboard": "Interactive HTML/JS Dashboard",
        "HTML Dashboard dashboard": "HTML Dashboard",
        "Week 1": "Day 1",
        "Week 2": "Day 2",
        "Week 3": "Day 3",
        "Week 4": "Day 4",
        "Week 5": "Day 5",
        "Week 6": "Day 6",
        "Week 7": "Day 7"
    }

    for file in files:
        if os.path.exists(file):
            update_document(file, replacements)
